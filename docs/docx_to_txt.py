# -*- coding: utf-8 -*-
"""
Convert DOCX to TXT - chỉ đánh dấu đáp án đúng.
Đáp án đúng (underline) → ★ prefix
Bỏ hết bold markers cho tiêu đề/đề mục.
"""
import sys, os
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def docx_to_txt(docx_path, txt_path):
    doc = Document(docx_path)
    lines = []
    
    for para in doc.paragraphs:
        if not para.text.strip():
            lines.append('')
            continue
        
        parts = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            
            is_underline = run.underline
            
            if is_underline:
                parts.append(f'★{text}')
            else:
                parts.append(text)
        
        line = ''.join(parts)
        lines.append(line)
    
    # Also extract tables
    for ti, table in enumerate(doc.tables):
        lines.append(f'\n--- TABLE {ti+1} ---')
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_parts = []
                for para in cell.paragraphs:
                    for run in para.runs:
                        text = run.text
                        if not text:
                            continue
                        if run.underline:
                            cell_parts.append(f'★{text}')
                        else:
                            cell_parts.append(text)
                cells.append(''.join(cell_parts))
            lines.append(' | '.join(cells))
        lines.append('--- END TABLE ---\n')
    
    with open(txt_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))
    
    print(f'✅ {os.path.basename(docx_path)} → {os.path.basename(txt_path)}')
    print(f'   {len(lines)} lines')

docs_dir = r'D:\Ontap\docs'

for f in os.listdir(docs_dir):
    if not f.endswith('.docx'):
        continue
    fl = f.lower()
    if 'sinh' in fl:
        out = 'sinh10.txt'
    elif 'ls 10' in fl or ('cương' in fl.lower()):
        out = 'su10.txt'
    elif 'lịch sử' in fl.lower() or 'lich su' in fl.lower():
        out = 'su12.txt'
    else:
        continue
    
    docx_to_txt(os.path.join(docs_dir, f), os.path.join(docs_dir, out))
